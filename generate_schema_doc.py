#!/usr/bin/env python
"""Generate Word document with DLT-META JSON schema documentation."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_table_borders(table):
    """Add borders to table."""
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def create_schema_table(doc, title, headers, rows):
    """Create a formatted table with schema information."""
    # Add title
    if title:
        heading = doc.add_heading(title, level=2)
        heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    
    # Create table
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    # Add borders
    add_table_borders(table)
    
    # Header row
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(255, 255, 255)
        # Set background color for header
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), '0070C0')
        hdr_cells[i]._element.get_or_add_tcPr().append(shading_elm)
    
    # Data rows
    for row_idx, row_data in enumerate(rows, start=1):
        row_cells = table.rows[row_idx].cells
        for col_idx, cell_data in enumerate(row_data):
            row_cells[col_idx].text = str(cell_data)
            for paragraph in row_cells[col_idx].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    
    doc.add_paragraph()  # Add spacing

def main():
    """Generate the Word document."""
    doc = Document()
    
    # Set document properties
    doc.core_properties.title = 'DLT-META JSON Configuration Schemas'
    doc.core_properties.author = 'DLT-META'
    
    # Title page
    title = doc.add_heading('DLT-META JSON Configuration Schemas', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    title.runs[0].font.size = Pt(24)
    
    subtitle = doc.add_paragraph('Complete Reference Guide')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.italic = True
    
    doc.add_paragraph()
    doc.add_paragraph(f'Generated: October 23, 2025')
    doc.add_page_break()
    
    # Table of Contents
    doc.add_heading('Table of Contents', level=1)
    toc = doc.add_paragraph()
    toc.add_run('1. Onboarding JSON Schema (Data Flow Specification)\n')
    toc.add_run('   1.1 Source Details Object\n')
    toc.add_run('   1.2 CDC Apply Changes Object\n')
    toc.add_run('   1.3 Apply Changes From Snapshot Object\n')
    toc.add_run('   1.4 Append Flows Array Object\n')
    toc.add_run('   1.5 Sink Object (DLT Sink API)\n')
    toc.add_run('2. Silver Transformations JSON Schema\n')
    toc.add_run('3. Data Quality Expectations (DQE) JSON Schema\n')
    doc.add_page_break()
    
    # Main Onboarding JSON Schema
    doc.add_heading('1. Onboarding JSON Schema (Data Flow Specification)', level=1)
    doc.add_paragraph('This is the main configuration file for defining data pipelines in DLT-META.')
    
    onboarding_data = [
        ['data_flow_id', 'String', '✅ Yes', 'Unique identifier for the pipeline'],
        ['data_flow_group', 'String', '✅ Yes', 'Group identifier for launching multiple pipelines under single DLT Pipeline'],
        ['source_system', 'String', '❌ Optional', 'Source system identifier (e.g., MYSQL, ORACLE, SAP)'],
        ['source_format', 'String', '✅ Yes', 'Source format: cloudFiles, eventhub, kafka, delta, snapshot'],
        ['source_details', 'Object', '✅ Yes', 'Source-specific configuration (see subsection 1.1)'],
        ['bronze_catalog_{env}', 'String', '⚠️ Recommended', 'Unity Catalog name for bronze layer (env = dev/prod/staging)'],
        ['bronze_database_{env}', 'String', '✅ Yes', 'Delta Lake bronze database name'],
        ['bronze_table', 'String', '✅ Yes', 'Delta Lake bronze table name'],
        ['bronze_table_comment', 'String', '❌ Optional', 'Bronze table comment/description'],
        ['bronze_reader_options', 'Object', '❌ Optional', 'Spark reader options (e.g., {"header": "true"})'],
        ['bronze_partition_columns', 'String/Array', '❌ Optional', 'Bronze table partition columns list'],
        ['bronze_table_cluster_by', 'Array', '❌ Optional', 'Bronze table cluster by columns list'],
        ['bronze_cdc_apply_changes', 'Object', '❌ Optional', 'Bronze CDC apply changes configuration (see subsection 1.2)'],
        ['bronze_apply_changes_from_snapshot', 'Object', '❌ Optional', 'Bronze apply changes from snapshot config (see subsection 1.3)'],
        ['bronze_table_path_{env}', 'String', '⚠️ Conditional', 'Bronze table storage path (mandatory if UC not enabled)'],
        ['bronze_table_properties', 'Object', '❌ Optional', 'DLT table properties (e.g., {"pipelines.autoOptimize.managed": "false"})'],
        ['bronze_sink', 'Object', '❌ Optional', 'DLT Sink API properties for bronze layer (see subsection 1.5)'],
        ['bronze_data_quality_expectations_json_{env}', 'String', '❌ Optional', 'Path to bronze DQE JSON file'],
        ['bronze_catalog_quarantine_{env}', 'String', '❌ Optional', 'Unity Catalog name for quarantine data'],
        ['bronze_database_quarantine_{env}', 'String', '❌ Optional', 'Database for quarantine data that fails expectations'],
        ['bronze_quarantine_table', 'String', '❌ Optional', 'Table name for quarantine data'],
        ['bronze_quarantine_table_comment', 'String', '❌ Optional', 'Quarantine table comment'],
        ['bronze_quarantine_table_path_{env}', 'String', '❌ Optional', 'Quarantine table storage path'],
        ['bronze_quarantine_table_partitions', 'Array', '❌ Optional', 'Quarantine table partition columns'],
        ['bronze_quarantine_table_cluster_by', 'Array', '❌ Optional', 'Quarantine table cluster columns'],
        ['bronze_quarantine_table_properties', 'Object', '❌ Optional', 'DLT table properties for quarantine table'],
        ['bronze_append_flows', 'Array', '❌ Optional', 'Bronze table append flows configuration (see subsection 1.4)'],
        ['silver_catalog_{env}', 'String', '⚠️ Recommended', 'Unity Catalog name for silver layer'],
        ['silver_database_{env}', 'String', '⚠️ Conditional', 'Silver database name (mandatory for silver layer)'],
        ['silver_table', 'String', '⚠️ Conditional', 'Silver table name (mandatory for silver layer)'],
        ['silver_table_comment', 'String', '❌ Optional', 'Silver table comment/description'],
        ['silver_partition_columns', 'Array', '❌ Optional', 'Silver table partition columns list'],
        ['silver_table_cluster_by', 'Array', '❌ Optional', 'Silver table cluster by columns list'],
        ['silver_cdc_apply_changes', 'Object', '❌ Optional', 'Silver CDC apply changes configuration (see subsection 1.2)'],
        ['silver_apply_changes_from_snapshot', 'Object', '❌ Optional', 'Silver apply changes from snapshot configuration'],
        ['silver_table_path_{env}', 'String', '⚠️ Conditional', 'Silver table storage path (mandatory if UC not enabled for silver layer)'],
        ['silver_table_properties', 'Object', '❌ Optional', 'DLT table properties for silver table'],
        ['silver_sink', 'Object', '❌ Optional', 'DLT Sink API properties for silver layer'],
        ['silver_transformation_json_{env}', 'String', '⚠️ Conditional', 'Path to silver transformation JSON file (mandatory for silver layer)'],
        ['silver_data_quality_expectations_json_{env}', 'String', '❌ Optional', 'Path to silver DQE JSON file'],
        ['silver_append_flows', 'Array', '❌ Optional', 'Silver table append flows configuration'],
    ]
    
    create_schema_table(doc, 'Main Configuration Fields', 
                       ['Field Name', 'Type', 'Mandatory', 'Description'],
                       onboarding_data)
    
    doc.add_page_break()
    
    # Source Details - CloudFiles
    doc.add_heading('1.1 Source Details Object', level=2)
    doc.add_heading('For CloudFiles Source Format:', level=3)
    
    cloudfiles_data = [
        ['source_database', 'String', '❌ Optional', 'Source database name'],
        ['source_table', 'String', '❌ Optional', 'Source table name'],
        ['source_path_{env}', 'String', '✅ Yes', 'Source file path for CloudFiles'],
        ['source_schema_path', 'String', '✅ Yes', 'Path to DDL schema file (Spark DDL format)'],
        ['source_metadata', 'Object', '❌ Optional', 'Metadata configuration for _metadata columns'],
        ['source_metadata.include_autoloader_metadata_column', 'String', '❌ Optional', '"True"/"False" to add _metadata column'],
        ['source_metadata.autoloader_metadata_col_name', 'String', '❌ Optional', 'Rename _metadata column (default: "source_metadata")'],
        ['source_metadata.select_metadata_cols', 'Object', '❌ Optional', 'Extract columns from _metadata (e.g., {"input_file_name": "_metadata.file_name"})'],
    ]
    
    create_schema_table(doc, '', 
                       ['Field Name', 'Type', 'Mandatory', 'Description'],
                       cloudfiles_data)
    
    # Source Details - EventHub
    doc.add_heading('For EventHub Source Format:', level=3)
    
    eventhub_data = [
        ['source_schema_path', 'String', '✅ Yes', 'Path to DDL schema file'],
        ['eventhub.accessKeyName', 'String', '✅ Yes', 'EventHub access key name'],
        ['eventhub.accessKeySecretName', 'String', '✅ Yes', 'EventHub access key secret name'],
        ['eventhub.name', 'String', '✅ Yes', 'EventHub name'],
        ['eventhub.secretsScopeName', 'String', '✅ Yes', 'Databricks secrets scope name'],
        ['eventhub.namespace', 'String', '✅ Yes', 'EventHub namespace'],
        ['eventhub.port', 'String', '✅ Yes', 'EventHub port (typically "9093")'],
        ['kafka.sasl.mechanism', 'String', '✅ Yes', 'SASL mechanism (typically "PLAIN")'],
        ['kafka.security.protocol', 'String', '✅ Yes', 'Security protocol (typically "SASL_SSL")'],
    ]
    
    create_schema_table(doc, '', 
                       ['Field Name', 'Type', 'Mandatory', 'Description'],
                       eventhub_data)
    
    # Source Details - Snapshot
    doc.add_heading('For Snapshot Source Format:', level=3)
    
    snapshot_data = [
        ['source_path_{env}', 'String', '✅ Yes', 'Path to snapshot files'],
        ['snapshot_format', 'String', '✅ Yes', 'Snapshot file format (e.g., "csv", "parquet")'],
    ]
    
    create_schema_table(doc, '', 
                       ['Field Name', 'Type', 'Mandatory', 'Description'],
                       snapshot_data)
    
    doc.add_page_break()
    
    # CDC Apply Changes
    doc.add_heading('1.2 CDC Apply Changes Object', level=2)
    
    cdc_data = [
        ['keys', 'Array', '✅ Yes', 'Primary key columns for CDC operations'],
        ['sequence_by', 'String', '✅ Yes', 'Column used for ordering changes (e.g., timestamp)'],
        ['scd_type', 'String', '✅ Yes', 'Slowly Changing Dimension type: "1" or "2"'],
        ['where', 'String', '❌ Optional', 'Filter condition for applying changes'],
        ['ignore_null_updates', 'Boolean', '❌ Optional', 'Whether to ignore null updates (default: false)'],
        ['apply_as_deletes', 'String', '❌ Optional', 'Expression to identify delete operations (e.g., "operation = \'DELETE\'")'],
        ['apply_as_truncates', 'String', '❌ Optional', 'Expression to identify truncate operations'],
        ['column_list', 'Array', '❌ Optional', 'Specific columns to include in CDC'],
        ['except_column_list', 'Array', '❌ Optional', 'Columns to exclude from CDC (e.g., ["operation", "_rescued_data"])'],
        ['track_history_column_list', 'Array', '❌ Optional', 'Columns to track history for (SCD Type 2)'],
        ['track_history_except_column_list', 'Array', '❌ Optional', 'Columns to exclude from history tracking'],
        ['flow_name', 'String', '❌ Optional', 'Custom flow name'],
        ['once', 'Boolean', '❌ Optional', 'Process data once (default: false)'],
        ['ignore_null_updates_column_list', 'Array', '❌ Optional', 'Specific columns to ignore null updates'],
        ['ignore_null_updates_except_column_list', 'Array', '❌ Optional', 'Columns to exclude from ignore null updates'],
    ]
    
    create_schema_table(doc, '', 
                       ['Field Name', 'Type', 'Mandatory', 'Description'],
                       cdc_data)
    
    doc.add_page_break()
    
    # Apply Changes From Snapshot
    doc.add_heading('1.3 Apply Changes From Snapshot Object', level=2)
    
    snapshot_changes_data = [
        ['keys', 'Array', '✅ Yes', 'Primary key columns for snapshot tracking'],
        ['scd_type', 'String', '✅ Yes', 'Slowly Changing Dimension type: "1" or "2"'],
        ['track_history_column_list', 'Array', '❌ Optional', 'Columns to track history for (SCD Type 2)'],
        ['track_history_except_column_list', 'Array', '❌ Optional', 'Columns to exclude from history tracking'],
    ]
    
    create_schema_table(doc, '', 
                       ['Field Name', 'Type', 'Mandatory', 'Description'],
                       snapshot_changes_data)
    
    # Append Flows
    doc.add_heading('1.4 Append Flows Array Object', level=2)
    
    append_flows_data = [
        ['name', 'String', '✅ Yes', 'Unique name for the append flow'],
        ['source_format', 'String', '✅ Yes', 'Source format for append flow'],
        ['create_streaming_table', 'Boolean', '✅ Yes', 'Whether to create streaming table (true/false)'],
        ['source_details', 'Object', '✅ Yes', 'Source details (same structure as main source_details)'],
        ['comment', 'String', '❌ Optional', 'Append flow comment/description'],
        ['reader_options', 'Object', '❌ Optional', 'Spark reader options'],
        ['spark_conf', 'Object', '❌ Optional', 'Spark configuration for this flow'],
        ['once', 'Boolean', '❌ Optional', 'Process data once (default: false)'],
        ['target', 'String', '❌ Optional', 'Target table name for silver append flows'],
    ]
    
    create_schema_table(doc, '', 
                       ['Field Name', 'Type', 'Mandatory', 'Description'],
                       append_flows_data)
    
    doc.add_page_break()
    
    # Sink Object
    doc.add_heading('1.5 Sink Object (DLT Sink API)', level=2)
    
    sink_data = [
        ['name', 'String', '✅ Yes', 'Sink name identifier'],
        ['format', 'String', '✅ Yes', 'Sink format: delta, kafka, eventhub'],
        ['options', 'Object', '✅ Yes', 'Format-specific options (e.g., {"tableName": "catalog.schema.table"} for delta)'],
        ['select_exp', 'Array', '❌ Optional', 'SQL expressions to apply before sinking'],
        ['where_clause', 'String', '❌ Optional', 'Filter condition before sinking'],
    ]
    
    create_schema_table(doc, '', 
                       ['Field Name', 'Type', 'Mandatory', 'Description'],
                       sink_data)
    
    doc.add_page_break()
    
    # Silver Transformations
    doc.add_heading('2. Silver Transformations JSON Schema', level=1)
    doc.add_paragraph('This file defines SQL transformations for silver layer tables.')
    
    silver_transform_data = [
        ['target_table', 'String', '✅ Yes', 'Target silver table name for transformation output'],
        ['select_exp', 'Array', '✅ Yes', 'Array of SQL expressions/column selections (e.g., ["concat(first_name,\' \',last_name) as full_name"])'],
        ['target_partition_cols', 'Array', '❌ Optional', 'Partition columns for the target table'],
        ['where_clause', 'Array', '❌ Optional', 'Array of filter conditions to apply (e.g., ["country = \'United States\'"])'],
    ]
    
    create_schema_table(doc, '', 
                       ['Field Name', 'Type', 'Mandatory', 'Description'],
                       silver_transform_data)
    
    # Example
    doc.add_heading('Example:', level=3)
    example = doc.add_paragraph(
        '[\n'
        '  {\n'
        '    "target_table": "cars_usa",\n'
        '    "select_exp": [\n'
        '      "concat(first_name,\' \',last_name) as full_name",\n'
        '      "country",\n'
        '      "brand",\n'
        '      "model"\n'
        '    ],\n'
        '    "where_clause": ["country = \'United States\'"]\n'
        '  }\n'
        ']'
    )
    example.runs[0].font.name = 'Courier New'
    example.runs[0].font.size = Pt(9)
    
    doc.add_page_break()
    
    # Data Quality Expectations
    doc.add_heading('3. Data Quality Expectations (DQE) JSON Schema', level=1)
    doc.add_paragraph('This file defines data quality rules for bronze or silver layers.')
    
    dqe_data = [
        ['expect', 'Object', '❌ Optional', 'Rules where failing records are INCLUDED in target dataset. Key=rule name, Value=SQL condition'],
        ['expect_or_fail', 'Object', '❌ Optional', 'Rules where failing records HALT pipeline execution. Key=rule name, Value=SQL condition'],
        ['expect_or_drop', 'Object', '❌ Optional', 'Rules where failing records are DROPPED from target dataset. Key=rule name, Value=SQL condition'],
        ['expect_or_quarantine', 'Object', '❌ Optional', 'Rules where failing records go to quarantine table (Bronze layer only). Key=rule name, Value=SQL condition'],
    ]
    
    create_schema_table(doc, '', 
                       ['Field Name', 'Type', 'Mandatory', 'Description'],
                       dqe_data)
    
    # Example
    doc.add_heading('Example:', level=3)
    dqe_example = doc.add_paragraph(
        '{\n'
        '  "expect_or_drop": {\n'
        '    "no_rescued_data": "_rescued_data IS NULL",\n'
        '    "valid_id": "id IS NOT NULL"\n'
        '  },\n'
        '  "expect_or_quarantine": {\n'
        '    "quarantine_rule": "_rescued_data IS NOT NULL OR id IS NULL OR amount=0"\n'
        '  }\n'
        '}'
    )
    dqe_example.runs[0].font.name = 'Courier New'
    dqe_example.runs[0].font.size = Pt(9)
    
    doc.add_page_break()
    
    # Notes
    doc.add_heading('Notes', level=1)
    notes = [
        '{env} placeholder represents environment suffix: _dev, _prod, _staging, _it, etc.',
        'Mandatory Fields (✅): Must be provided in all cases',
        'Conditional Fields (⚠️ Conditional): Required based on context (e.g., bronze_table_path_{env} is mandatory only when Unity Catalog is not enabled; silver fields are mandatory when defining silver layer)',
        'Recommended Fields (⚠️ Recommended): Highly recommended for Unity Catalog environments',
        'Optional Fields (❌): Can be omitted',
        'All JSON files must be valid JSON format and accessible from the Databricks workspace',
        'Schema files referenced in source_schema_path should be in Spark DDL format',
        'Source code validation occurs in src/onboard_dataflowspec.py and src/dataflow_spec.py'
    ]
    
    for note in notes:
        p = doc.add_paragraph(note, style='List Bullet')
        p.runs[0].font.size = Pt(10)
    
    # Save document
    output_path = '/Users/paulambasu/Downloads/Development/Databricks/dlt-meta/DLT-META_JSON_Schema_Documentation.docx'
    doc.save(output_path)
    print(f'Document saved to: {output_path}')

if __name__ == '__main__':
    main()
